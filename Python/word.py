from docx import Document
from docx.shared import Inches

document=Document()
document.add_heading("Document Title",0)

p=document.add_paragraph('A plain papagraph having some')
p.add_run('blod').blod=True
p.add_run('add some')
p.add_run('italic.').italic=True

document.add_heading('Heading level 1',level=1)
document.add_paragraph("Intense quote",style='IntenseQuote')

document.add_paragraph('firdt item in unorderd list',style='ListBuilet')
document.add_paragraph('hellow everybody',style='ListBuilet')

#document.add_picture('a.jpg',width=Inches(1.25))

table=document.add_table(rows=1,cols=3)
hdr_cells=table.rows[0].cells
hdr_cells[0].text='Qty'
hdr_cells[1].text='Id'
hdr_cells[2].text='desc'

for itemd in recordset:
    row_cells=table.add_rows().cells
    row_cells[0].text=str(item.qty)
    row_cells[1].text=str(item.id)
    row_cells[2].text=item.desc

document.add_page_break()
document.save('demo.doc')

